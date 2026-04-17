import base64
import io
from docx import Document


def parse_docx(docx_bytes: bytes) -> dict:
    """
    Returns {"text": str, "images": [{"data": base64_str, "mime_type": str, "page": None, "index": int}]}
    Raises ValueError if no extractable text is found.
    Note: nested tables (tables inside table cells) are not extracted.
    """
    doc = Document(io.BytesIO(docx_bytes))

    # 1. Extract paragraphs
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # 2. Extract table cells (skip merged cell duplicates)
    seen_cell_ids = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if id(cell._tc) not in seen_cell_ids and cell.text.strip():
                    seen_cell_ids.add(id(cell._tc))
                    paragraphs.append(cell.text.strip())

    full_text = "\n".join(paragraphs).strip()

    # 3. Extract images (guard against externally-linked images)
    images = []
    img_index = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype and not rel.is_external:
            image_bytes = rel.target_part.blob
            mime_type = rel.target_part.content_type
            images.append({
                "data": base64.b64encode(image_bytes).decode("utf-8"),
                "mime_type": mime_type,
                "page": None,
                "index": img_index,
            })
            img_index += 1

    # 4. Validate text after image extraction
    if not full_text:
        raise ValueError("DOCX contains no extractable text — the document may be empty or image-only.")

    # 5. Return result
    return {"text": full_text, "images": images}
