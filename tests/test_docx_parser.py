import io
import pytest
from docx import Document
from parsers.docx import parse_docx


def make_test_docx(text: str) -> bytes:
    doc = Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_parse_docx_returns_text():
    docx_bytes = make_test_docx("Deliverable: API Gateway Integration")
    result = parse_docx(docx_bytes)
    assert "API Gateway Integration" in result["text"]


def test_parse_docx_returns_images_list():
    docx_bytes = make_test_docx("Some content")
    result = parse_docx(docx_bytes)
    assert "images" in result
    assert isinstance(result["images"], list)


def test_parse_docx_no_images_returns_empty_list():
    docx_bytes = make_test_docx("Some content")
    result = parse_docx(docx_bytes)
    assert result["images"] == []


def test_parse_docx_raises_on_empty_text():
    doc = Document()
    buf = io.BytesIO()
    doc.save(buf)
    with pytest.raises(ValueError, match="DOCX contains no extractable text"):
        parse_docx(buf.getvalue())


def test_parse_docx_extracts_table_cell_text():
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Phase"
    table.cell(0, 1).text = "Build"
    buf = io.BytesIO()
    doc.save(buf)
    result = parse_docx(buf.getvalue())
    assert "Phase" in result["text"]
    assert "Build" in result["text"]


def test_parse_docx_image_has_required_keys():
    """Image dicts must have data, mime_type, page, and index keys.
    Uses a 200x200 image so it passes the size filter."""
    png_bytes = make_png(200, 200)
    doc = Document()
    doc.add_paragraph("Some content")
    doc.add_picture(io.BytesIO(png_bytes))
    buf = io.BytesIO()
    doc.save(buf)
    result = parse_docx(buf.getvalue())
    assert len(result["images"]) >= 1
    img = result["images"][0]
    assert "data" in img
    assert "mime_type" in img
    assert "page" in img
    assert "index" in img


def make_png(w: int, h: int) -> bytes:
    """Create a minimal valid PNG of given pixel dimensions."""
    import struct, zlib
    sig = b'\x89PNG\r\n\x1a\n'
    ihdr_data = struct.pack('>IIBBBBB', w, h, 8, 2, 0, 0, 0)
    ihdr = struct.pack('>I', 13) + b'IHDR' + ihdr_data
    ihdr += struct.pack('>I', zlib.crc32(b'IHDR' + ihdr_data) & 0xffffffff)
    raw = (b'\x00' + b'\xff\xff\xff' * w) * h
    idat_data = zlib.compress(raw)
    idat = struct.pack('>I', len(idat_data)) + b'IDAT' + idat_data
    idat += struct.pack('>I', zlib.crc32(b'IDAT' + idat_data) & 0xffffffff)
    iend = b'\x00\x00\x00\x00IEND\xaeB`\x82'
    return sig + ihdr + idat + iend


def test_parse_docx_skips_images_smaller_than_200x200():
    """Images below 200x200 pixels must be filtered from the returned list."""
    png_bytes = make_png(10, 10)
    doc = Document()
    doc.add_paragraph("Some content")
    doc.add_picture(io.BytesIO(png_bytes))
    buf = io.BytesIO()
    doc.save(buf)
    result = parse_docx(buf.getvalue())
    assert result["images"] == [], "tiny image (10x10) should be filtered out"


def test_parse_docx_keeps_images_at_least_200x200():
    """Images that meet the 200x200 threshold must be kept."""
    png_bytes = make_png(200, 200)
    doc = Document()
    doc.add_paragraph("Some content")
    doc.add_picture(io.BytesIO(png_bytes))
    buf = io.BytesIO()
    doc.save(buf)
    result = parse_docx(buf.getvalue())
    assert len(result["images"]) >= 1, "200x200 image should NOT be filtered"
